"""
Document Processing Utilities

This module provides utilities for extracting text from various document formats
including PDF, Word documents, and plain text files.
"""

import io

# Removed direct logging import - using unified config

# Import document processing libraries with availability checks
try:
    import PyPDF2

    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import pdfplumber

    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ..config.logfire_config import get_logger, logfire

logger = get_logger(__name__)


def extract_text_from_document(file_content: bytes, filename: str, content_type: str) -> str:
    """
    Extract text from various document formats.

    Args:
        file_content: Raw file bytes
        filename: Name of the file
        content_type: MIME type of the file

    Returns:
        Extracted text content

    Raises:
        ValueError: If the file format is not supported
        Exception: If extraction fails
    """
    try:
        # PDF files
        if content_type == "application/pdf" or filename.lower().endswith(".pdf"):
            return extract_text_from_pdf(file_content)

        # Word documents
        elif content_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ] or filename.lower().endswith((".docx", ".doc")):
            return extract_text_from_docx(file_content)

        # Text files (markdown, txt, etc.)
        elif content_type.startswith("text/") or filename.lower().endswith((
            ".txt",
            ".md",
            ".markdown",
            ".rst",
        )):
            return file_content.decode("utf-8", errors="ignore")

        else:
            raise ValueError(f"Unsupported file format: {content_type} ({filename})")

    except Exception as e:
        logfire.error(
            "Document text extraction failed",
            filename=filename,
            content_type=content_type,
            error=str(e),
        )
        raise Exception(f"Failed to extract text from {filename}: {str(e)}")


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF using both PyPDF2 and pdfplumber for best results.

    Args:
        file_content: Raw PDF bytes

    Returns:
        Extracted text content
    """
    if not PDFPLUMBER_AVAILABLE and not PYPDF2_AVAILABLE:
        raise Exception(
            "No PDF processing libraries available. Please install pdfplumber and PyPDF2."
        )

    text_content = []

    # First try with pdfplumber (better for complex layouts)
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        logfire.warning(f"pdfplumber failed on page {page_num + 1}: {e}")
                        continue

            # If pdfplumber got good results, use them
            if text_content and len("\n".join(text_content).strip()) > 100:
                return "\n\n".join(text_content)

        except Exception as e:
            logfire.warning(f"pdfplumber extraction failed: {e}, trying PyPDF2")

    # Fallback to PyPDF2
    if PYPDF2_AVAILABLE:
        try:
            text_content = []
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logfire.warning(f"PyPDF2 failed on page {page_num + 1}: {e}")
                    continue

            if text_content:
                return "\n\n".join(text_content)
            else:
                raise Exception("No text could be extracted from PDF")

        except Exception as e:
            raise Exception(f"PyPDF2 failed to extract text: {str(e)}")

    # If we get here, no libraries worked
    raise Exception("Failed to extract text from PDF - no working PDF libraries available")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from Word documents (.docx).

    Args:
        file_content: Raw DOCX bytes

    Returns:
        Extracted text content
    """
    if not DOCX_AVAILABLE:
        raise Exception("python-docx library not available. Please install python-docx.")

    try:
        doc = DocxDocument(io.BytesIO(file_content))
        text_content = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))

        if not text_content:
            raise Exception("No text content found in document")

        return "\n\n".join(text_content)

    except Exception as e:
        raise Exception(f"Failed to extract text from Word document: {str(e)}")
